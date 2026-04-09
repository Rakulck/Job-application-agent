import os
import re
import json
import tempfile
from io import BytesIO
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from supabase import create_client
from dotenv import load_dotenv

load_dotenv()

SUPABASE_URL         = os.getenv("SUPABASE_URL")
SUPABASE_KEY         = os.getenv("SUPABASE_KEY")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

supabase         = create_client(SUPABASE_URL, SUPABASE_KEY)
supabase_admin   = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)  # bypasses RLS

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Colours / fonts (hardcoded — Groq never touches these)
# ---------------------------------------------------------------------------
COLOR_NAME      = RGBColor(0x00, 0x00, 0x00)   # black
COLOR_HEADING   = RGBColor(0x00, 0x70, 0xC0)   # blue #0070C0
COLOR_BODY      = RGBColor(0x00, 0x00, 0x00)   # black
COLOR_LIGHT     = RGBColor(0x00, 0x00, 0x00)   # black
FONT_MAIN       = "Times New Roman"

# Skill category mappings — builder groups flat skill list into these
SKILL_CATEGORIES = {
    "Programming Languages":    ["JavaScript", "TypeScript", "Python", "C++", "Java", "SQL", "Dart"],
    "Frameworks & Libraries":   ["React.js", "React Native", "Next.js", "Node.js", "Vue", "Angular",
                                  "Flutter", "Express", "Tailwind CSS", "HTML", "CSS", "Redux", "Expo"],
    "Databases & Storage":      ["MySQL", "MongoDB", "Firebase", "Supabase", "NoSQL", "PostgreSQL",
                                  "Redis", "DynamoDB"],
    "Cloud & DevOps Tools":     ["AWS", "Git", "GitHub", "Docker", "Figma", "Android Studio",
                                  "Kubernetes", "GCP", "Azure", "Canva", "Cursor", "V0", "Bolt"],
    "Testing Frameworks":       ["Jest", "React Testing Library", "Unit Testing", "Integration Testing",
                                  "Detox", "Cypress"],
    "AI & API Integrations":    ["Gemini API", "OpenAI API", "Spoonacular API", "REST APIs", "GraphQL"],
}


def _categorize_skills(flat_skills: list) -> dict:
    """Group a flat skills list into categories. Uncategorized go to 'Other'."""
    used = set()
    result = {}
    for cat, keywords in SKILL_CATEGORIES.items():
        matched = [s for s in flat_skills if s in keywords]
        if matched:
            result[cat] = matched
            used.update(matched)
    other = [s for s in flat_skills if s not in used]
    if other:
        result["Other"] = other
    return result


# ===========================================================================
# python-docx builder
# ===========================================================================

def _add_run(para, text, bold=False, italic=False, font_size=10,
             color=None, font_name=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name or FONT_MAIN
    if color:
        run.font.color.rgb = color
    return run


def _add_hyperlink(para, text, url, font_size=11.5):
    """Add a clickable hyperlink run to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue underline style
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0070C0")
    rPr.append(color_el)

    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz_el)

    szCs_el = OxmlElement("w:szCs")
    szCs_el.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(szCs_el)

    fonts_el = OxmlElement("w:rFonts")
    fonts_el.set(qn("w:ascii"), FONT_MAIN)
    fonts_el.set(qn("w:hAnsi"), FONT_MAIN)
    rPr.append(fonts_el)

    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)


def _add_hr(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0070C0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _para(doc, text="", bold=False, italic=False, font_size=10,
          color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        _add_run(p, text, bold=bold, italic=italic,
                 font_size=font_size, color=color)
    return p


JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
CENTER  = WD_ALIGN_PARAGRAPH.CENTER
TAB_SKILLS  = Inches(2.5)   # skill category tab stop
TAB_DATE    = Inches(7.5)   # right-align tab for dates

# ── Template font sizes (locked from approved doc 173438) ──
FS_NAME     = 15.0   # name header
FS_HEADING  = 13.0   # section headings (blue)
FS_EXP_CO   = 12.5   # experience company line
FS_BODY     = 11.5   # everything else: contact, summary, skills, bullets, education


def _section_heading(doc, title: str):
    """13pt bold blue ALL CAPS heading with bottom border (HR) baked in."""
    p = doc.add_paragraph()
    p.alignment = JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    _add_run(p, title.upper() + ":", bold=True, font_size=FS_HEADING, color=COLOR_HEADING)

    # Bottom border directly on the heading paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0070C0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


MAX_EXPERIENCE_ENTRIES = 3   # hard cap — keeps resume to one page
MAX_BULLETS_PER_JOB   = 3   # hard cap — keeps resume to one page


def build_docx(data: dict, company: str = "") -> str:
    """Build a .docx resume from tailored JSON. Returns file path."""
    doc = Document()

    # Page: Letter size, 0 top/bottom, 0.51cm left/right (matches reference)
    for sec in doc.sections:
        sec.page_width    = Cm(21.59)
        sec.page_height   = Cm(27.94)
        sec.top_margin    = Cm(1.27)
        sec.bottom_margin = Cm(1.27)
        sec.left_margin   = Cm(1.27)
        sec.right_margin  = Cm(1.27)

    # ---- Name: 14pt bold centered ----
    name_p = doc.add_paragraph()
    name_p.alignment = CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after  = Pt(0)
    _add_run(name_p, data.get("name", ""), bold=True, font_size=FS_NAME)

    # ---- Contact: plain | hyperlinks, 10.5pt centered ----
    contact_p = doc.add_paragraph()
    contact_p.alignment = CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after  = Pt(0)

    plain_parts = [data[f] for f in ["location", "email", "phone"] if data.get(f)]
    _add_run(contact_p, "  |  ".join(plain_parts), font_size=FS_BODY)

    link_labels = {"linkedin": "LinkedIn", "portfolio": "Portfolio", "github": "GitHub"}
    for field in ["linkedin", "portfolio", "github"]:
        val = data.get(field, "")
        if not val:
            continue
        _add_run(contact_p, "  |  ", font_size=FS_BODY)
        url = val if val.startswith("http") else f"https://{val}"
        _add_hyperlink(contact_p, link_labels[field], url, font_size=FS_BODY)

    # ---- Professional Summary ----
    _section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.alignment = JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _add_run(p, data.get("summary", ""), font_size=FS_BODY)

    # ---- Technical Skills: "Category:\t values" per line ----
    skills = data.get("skills", [])
    if skills:
        _section_heading(doc, "Technical Skills")
        categorized = _categorize_skills(skills)
        for cat, items in categorized.items():
            p = doc.add_paragraph()
            p.alignment = JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.tab_stops.add_tab_stop(TAB_SKILLS)
            _add_run(p, f"{cat}:", bold=True, font_size=FS_BODY)
            _add_run(p, f"\t{',  '.join(items)}", bold=False, font_size=FS_BODY)

    # ---- Professional Experience ----
    experience = data.get("experience", [])[:MAX_EXPERIENCE_ENTRIES]
    if experience:
        _section_heading(doc, "Professional Experience")
        for exp in experience:
            exp_company = exp.get("company", "")
            location    = exp.get("location", "")
            title       = exp.get("title", "")
            dates       = f"{exp.get('start_date', '')} \u2013 {exp.get('end_date', '')}"

            # One line: "Company | Title \t Dates" — 12.5pt bold
            p = doc.add_paragraph()
            p.alignment = JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.tab_stops.add_tab_stop(TAB_DATE, WD_ALIGN_PARAGRAPH.RIGHT)
            _add_run(p, f"{exp_company} | {title}", bold=True, font_size=FS_EXP_CO)
            r = p.add_run(f"\t{dates}")
            r.bold = True
            r.font.size = Pt(FS_EXP_CO)
            r.font.name = FONT_MAIN

            # Bullets: List Paragraph, 10.5pt, justified (capped for one-page)
            for bullet in exp.get("bullets", [])[:MAX_BULLETS_PER_JOB]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.alignment = JUSTIFY
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after  = Pt(0)
                _add_run(bp, bullet, font_size=FS_BODY)

    # ---- Education ----
    education = data.get("education", [])
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            p = doc.add_paragraph()
            p.alignment = JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            _add_run(p, edu.get("degree", ""), bold=True, font_size=FS_BODY)
            rest = f" - {edu.get('school', '')}, {edu.get('location', '')}"
            if edu.get("graduation"):
                rest += f",  {edu.get('graduation', '')}"
            _add_run(p, rest, bold=False, font_size=FS_BODY)

    # Filename: FirstLast_Company.pdf (first + last name only, no middle initial)
    name_parts   = data.get("name", "resume").split()
    clean_name   = f"{name_parts[0]}_{name_parts[-1]}" if len(name_parts) >= 2 else "_".join(name_parts)
    company_slug = re.sub(r"[^A-Za-z0-9]+", "_", (company or "")).strip("_")
    slug = f"{clean_name}_{company_slug}" if company_slug else clean_name
    path = os.path.join(OUTPUT_DIR, f"{slug}.docx")
    doc.save(path)
    print(f"[builder] .docx saved: {path}")
    return path


# ===========================================================================
# PDF via docx2pdf (uses Microsoft Word — preserves exact DOCX formatting)
# ===========================================================================

def build_pdf(data: dict, docx_path: str) -> str:
    """Convert DOCX to PDF using Word. Returns PDF file path."""
    from docx2pdf import convert
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    print(f"[builder] PDF saved: {pdf_path}")
    return pdf_path


# ===========================================================================
# Supabase Storage upload
# ===========================================================================

def upload_pdf(pdf_path: str, job_id: str) -> str:
    """Upload PDF to Supabase Storage 'resumes' bucket. Returns public URL."""
    filename = os.path.basename(pdf_path)
    with open(pdf_path, "rb") as f:
        data = f.read()

    supabase_admin.storage.from_("resumes").upload(
        path=filename,
        file=data,
        file_options={"content-type": "application/pdf", "upsert": "true"},
    )
    base_url = supabase_admin.storage.from_("resumes").get_public_url(filename)
    # Bust CDN cache so the browser always fetches the latest version
    ts = datetime.now().strftime("%Y%m%d%H%M%S")
    url = f"{base_url}?v={ts}"
    print(f"[builder] uploaded to Supabase Storage: {url}")
    return url


# ===========================================================================
# Main pipeline
# ===========================================================================

def run_builder(job_id: str = None, force: bool = False):
    """Build PDFs for resumes. force=True rebuilds even if pdf_url already set."""
    if job_id:
        rows = supabase.table("resumes").select("*").eq("job_id", job_id).execute().data
    else:
        all_rows = supabase.table("resumes").select("*").execute().data
        rows = all_rows if force else [r for r in all_rows if not r.get("pdf_url")]

    if not rows:
        print("[builder] nothing to build")
        return

    print(f"[builder] building {len(rows)} resumes")

    # Pre-fetch all job companies in one query
    job_companies = {
        r["job_id"]: r.get("company", "")
        for r in supabase.table("jobs").select("job_id, company").execute().data
    }

    for row in rows:
        jid = row["job_id"]
        tailored = row["tailored_json"]
        company  = job_companies.get(jid, "")
        print(f"[builder] processing job_id: {jid} ({company})")

        try:
            docx_path = build_docx(tailored, company=company)
            pdf_path = build_pdf(tailored, docx_path)
            url = upload_pdf(pdf_path, jid)

            supabase.table("resumes").update({"pdf_url": url}).eq("id", row["id"]).execute()
            print(f"[builder] done: {jid}")

        except Exception as e:
            print(f"[builder] error for {jid}: {e}")

    print("[builder] all done")


if __name__ == "__main__":
    import sys
    jid   = sys.argv[1] if len(sys.argv) > 1 and sys.argv[1] != "--force" else None
    force = "--force" in sys.argv
    run_builder(job_id=jid, force=force)
